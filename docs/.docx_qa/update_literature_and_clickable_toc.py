from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DOCX = Path("DB2_PROJECT_REPORT_template_style.docx")
TABLE_HELPERS = Path(
    r"C:\Users\marve\.codex\plugins\cache\openai-primary-runtime\documents\26.430.10722"
    r"\skills\documents\scripts"
)
sys.path.insert(0, str(TABLE_HELPERS))
from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa  # noqa: E402


HEADING_BLUE = RGBColor(0x0F, 0x47, 0x61)
TOC_BLUE = RGBColor(0x21, 0x5E, 0x99)
LIGHT_BLUE = "EAF2F8"

USED_LITERATURE = [
    (
        "Security Industry Association",
        "Open Supervised Device Protocol (OSDP).",
        "https://www.securityindustry.org/industry-standards/open-supervised-device-protocol/",
    ),
    (
        "MongoDB Documentation",
        "MongoDB Manual: data modeling, CRUD operations, and indexing.",
        "https://www.mongodb.com/docs/manual/",
    ),
    (
        "MongoDB Documentation",
        "PyMongo driver documentation: MongoClient, databases, collections, and collection operations.",
        "https://www.mongodb.com/docs/languages/python/pymongo-driver/current/",
    ),
    (
        "Pallets Projects",
        "Flask documentation for Python web application structure and routing.",
        "https://flask.palletsprojects.com/",
    ),
    (
        "Flask-SocketIO",
        "Flask-SocketIO documentation for real-time event transport in Flask applications.",
        "https://flask-socketio.readthedocs.io/",
    ),
    (
        "React",
        "React documentation for component-based frontend development.",
        "https://react.dev/",
    ),
    (
        "Docker Documentation",
        "Docker run and container documentation used for MongoDB deployment notes.",
        "https://docs.docker.com/engine/containers/run/",
    ),
    (
        "Raspberry Pi Documentation",
        "Raspberry Pi documentation for Linux-based controller host context.",
        "https://www.raspberrypi.com/documentation/",
    ),
    (
        "STMicroelectronics",
        "STM32F103C8 product and datasheet material for the STM32 Blue Pill bridge MCU.",
        "https://www.st.com/en/microcontrollers-microprocessors/stm32f103c8.html",
    ),
    (
        "Analog Devices / Maxim Integrated",
        "MAX481/MAX483/MAX485/MAX487-MAX491 RS-485/RS-422 transceiver datasheet.",
        "https://www.analog.com/media/en/technical-documentation/data-sheets/MAX1487-MAX491.pdf",
    ),
]


def set_run_font(run, *, name="Times New Roman", size=11, bold=None, italic=None, color=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_paragraph_text(paragraph, text: str, *, size=11, bold=None, color=None) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def format_cell(cell, *, header=False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if header:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), LIGHT_BLUE)
    for paragraph in cell.paragraphs:
        paragraph.style = "Normal"
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            set_run_font(run, size=9, bold=True if header else None)


def replace_suggested_ui_test_pass(doc: Document) -> None:
    heading_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "G.2.14 Suggested UI Test Pass":
            heading_idx = idx
            break
    if heading_idx is None:
        if any(
            p.text.strip() == "G.2.14 Used Literature" and p.style.name == "Heading 3"
            for p in doc.paragraphs
        ):
            return
        raise RuntimeError("Could not find G.2.14 Suggested UI Test Pass")

    heading = doc.paragraphs[heading_idx]
    set_paragraph_text(heading, "G.2.14 Used Literature", size=14, bold=True, color=HEADING_BLUE)

    # Remove the five old test-pass paragraphs immediately after the heading.
    for paragraph in list(doc.paragraphs[heading_idx + 1 : heading_idx + 6]):
        delete_paragraph(paragraph)

    intro = doc.add_paragraph(
        "The following literature and documentation were used as technical background for the database, "
        "backend, frontend, embedded bridge, and deployment parts of the project."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.paragraph_format.space_after = Pt(6)
    for run in intro.runs:
        set_run_font(run, size=11)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Source", "Used For", "Reference"]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
        format_cell(table.rows[0].cells[i], header=True)

    for source, use, url in USED_LITERATURE:
        row = table.add_row()
        values = [source, use, url]
        for i, text in enumerate(values):
            row.cells[i].text = text
            format_cell(row.cells[i])

    content_width = section_content_width_dxa(doc.sections[0])
    widths = column_widths_from_weights([1.45, 3.25, 3.3], total_width_dxa=content_width)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=content_width,
        indent_dxa=0,
        cell_margins_dxa={"top": 95, "bottom": 95, "start": 120, "end": 120},
    )


def normalize_heading_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def bookmark_name(index: int) -> str:
    return f"toc_target_{index:03d}"


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    p = paragraph._p
    # Avoid duplicate bookmarks on reruns.
    for old in list(p.findall(qn("w:bookmarkStart"))):
        p.remove(old)
    for old in list(p.findall(qn("w:bookmarkEnd"))):
        p.remove(old)

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    insert_at = 1 if len(p) and p[0].tag == qn("w:pPr") else 0
    p.insert(insert_at, start)
    p.append(end)


def make_internal_link(paragraph, text: str, anchor: str, *, level: int) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Pt(20 if level == 2 else 38 if level == 3 else 0)
    paragraph.paragraph_format.space_after = Pt(1)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "Times New Roman")
    r_pr.append(fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22" if level == 1 else "20")
    r_pr.append(size)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "215E99" if level == 1 else "000000")
    r_pr.append(color)
    if level == 1:
        bold = OxmlElement("w:b")
        r_pr.append(bold)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(r_pr)
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def visible_toc_range(doc: Document) -> range:
    start = end = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "TABLE OF CONTENTS":
            start = idx + 1
            continue
        if start is not None and paragraph.style.name == "Heading 1":
            end = idx
            break
    if start is None or end is None:
        raise RuntimeError("Could not determine visible TOC range")
    return range(start, end)


def heading_level(style_name: str) -> int | None:
    if style_name == "Heading 1":
        return 1
    if style_name == "Heading 2":
        return 2
    if style_name == "Heading 3":
        return 3
    return None


def add_used_literature_to_toc(doc: Document) -> None:
    toc_range = visible_toc_range(doc)
    first_body_heading = toc_range.stop

    # Remove any earlier misplaced TOC entry, such as one inserted after the TOC page break.
    for paragraph in list(doc.paragraphs[toc_range.start : first_body_heading]):
        if (
            paragraph.text.strip() == "G.2.14 Used Literature"
            and paragraph.style.name == "Normal"
        ):
            delete_paragraph(paragraph)

    toc_range = visible_toc_range(doc)
    first_body_heading = toc_range.stop
    entries = [normalize_heading_text(doc.paragraphs[i].text) for i in toc_range if doc.paragraphs[i].text.strip()]
    if "G.2.14 Used Literature" in entries:
        return

    insert_before = doc.paragraphs[first_body_heading]
    for idx in toc_range:
        if "w:br" in doc.paragraphs[idx]._p.xml:
            insert_before = doc.paragraphs[idx]
            break
    p = insert_before.insert_paragraph_before("G.2.14 Used Literature")
    p.paragraph_format.left_indent = Pt(38)
    p.paragraph_format.space_after = Pt(1)
    for run in p.runs:
        set_run_font(run, size=10)


def make_toc_clickable(doc: Document) -> None:
    add_used_literature_to_toc(doc)

    headings: dict[str, tuple[str, int]] = {}
    bookmark_id = 1
    for paragraph in doc.paragraphs:
        level = heading_level(paragraph.style.name)
        if level is None:
            continue
        text = normalize_heading_text(paragraph.text)
        if not text or text in headings:
            continue
        name = bookmark_name(bookmark_id)
        add_bookmark(paragraph, name, bookmark_id)
        headings[text] = (name, level)
        bookmark_id += 1

    for idx in visible_toc_range(doc):
        paragraph = doc.paragraphs[idx]
        text = normalize_heading_text(paragraph.text)
        if not text or text not in headings:
            continue
        anchor, level = headings[text]
        make_internal_link(paragraph, text, anchor, level=level)


def main() -> None:
    doc = Document(DOCX)
    replace_suggested_ui_test_pass(doc)
    make_toc_clickable(doc)
    doc.save(DOCX)
    print(DOCX.resolve())


if __name__ == "__main__":
    main()
