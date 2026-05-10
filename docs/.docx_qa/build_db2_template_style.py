from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parents[1]
SOURCE_DOCX = DOCS_DIR / "DB2_PROJECT_REPORT_latest.docx"
REFERENCE_DOCX = DOCS_DIR / "Databases I Project template Music Streaming M.Velichkovski.docx"
OUTPUT_DOCX = DOCS_DIR / "DB2_PROJECT_REPORT_template_style.docx"
BUILD_DIR = DOCS_DIR / ".docx_qa" / "build"
PANDOC_OUT = BUILD_DIR / "db2_reference_converted.docx"

TABLE_HELPERS = Path(
    r"C:\Users\marve\.codex\plugins\cache\openai-primary-runtime\documents\26.430.10722"
    r"\skills\documents\scripts"
)
sys.path.insert(0, str(TABLE_HELPERS))
from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa  # noqa: E402


BLUE = RGBColor(0x21, 0x5E, 0x99)
HEADING_BLUE = RGBColor(0x0F, 0x47, 0x61)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F2F2"


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_run_font(run, *, name="Times New Roman", size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, *, name="Times New Roman", size=11, bold=None, italic=None, color=None) -> None:
    for run in paragraph.runs:
        set_run_font(run, name=name, size=size, bold=bold, italic=italic, color=color)


def add_before(anchor, text="", *, style=None, alignment=None, size=11, bold=None, italic=None,
               color=None, left_indent=None, space_after=None):
    paragraph = anchor.insert_paragraph_before("", style=style)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    if alignment is not None:
        paragraph.alignment = alignment
    if left_indent is not None:
        paragraph.paragraph_format.left_indent = Inches(left_indent)
    if space_after is not None:
        paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph


def add_page_break_before(anchor):
    paragraph = anchor.insert_paragraph_before("")
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    return paragraph


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def has_numbering(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return p_pr is not None and p_pr.numPr is not None


def run_pandoc() -> None:
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("pandoc was not found on PATH")
    cmd = [
        pandoc,
        str(SOURCE_DOCX),
        "--reference-doc",
        str(REFERENCE_DOCX),
        "-o",
        str(PANDOC_OUT),
    ]
    subprocess.run(cmd, check=True, cwd=DOCS_DIR)


def configure_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.start_type = WD_SECTION_START.NEW_PAGE

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = HEADING_BLUE
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(16)
    styles["Heading 2"].font.color.rgb = HEADING_BLUE
    styles["Heading 3"].font.name = "Times New Roman"
    styles["Heading 3"].font.size = Pt(14)
    styles["Heading 3"].font.color.rgb = HEADING_BLUE


def collect_toc_entries(doc: Document) -> list[tuple[int, str]]:
    entries: list[tuple[int, str]] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name
        if style == "Heading 1":
            entries.append((1, text))
        elif style == "Heading 2":
            entries.append((2, text))
    return entries


def insert_front_matter(doc: Document, toc_entries: list[tuple[int, str]]) -> None:
    for paragraph in list(doc.paragraphs[:6]):
        delete_paragraph(paragraph)

    anchor = doc.paragraphs[0]

    add_before(anchor, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_before(anchor, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_before(
        anchor,
        "Databases II Course Project",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=20,
        space_after=10,
    )
    add_before(
        anchor,
        "OSDP Access Controller Proof of Concept Database",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size=20,
        space_after=18,
    )
    add_before(anchor, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_before(anchor, "Team Members:", size=14, space_after=4)
    add_before(anchor, "Student Name: Martin Velichkovski (Faculty: CNS)", style="List Paragraph", size=14, space_after=4)
    add_before(anchor, "Mentor: Dijana Capeska Bogatinoska, Ph.D.", size=14, space_after=4)
    add_before(anchor, "Date: May 2026", size=14, bold=True, space_after=4)
    add_page_break_before(anchor)

    add_before(anchor, "TABLE OF CONTENTS", size=16, bold=True, color=BLUE, space_after=12)
    for level, text in toc_entries:
        indent = 0.28 if level == 2 else 0.0
        size = 10 if level == 2 else 11
        add_before(
            anchor,
            text,
            size=size,
            bold=(level == 1),
            color=HEADING_BLUE if level == 1 else None,
            left_indent=indent,
            space_after=1,
        )
    add_page_break_before(anchor)


def format_paragraphs(doc: Document) -> None:
    previous_had_graphic = False
    for paragraph in doc.paragraphs:
        style = paragraph.style.name
        text = paragraph.text.strip()
        has_graphic = "graphic" in paragraph._p.xml

        if style == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            set_paragraph_font(paragraph, size=18, bold=True, color=HEADING_BLUE)
        elif style == "Heading 2":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(4)
            set_paragraph_font(paragraph, size=16, bold=True, color=HEADING_BLUE)
        elif style == "Heading 3":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(3)
            set_paragraph_font(paragraph, size=14, bold=True, color=HEADING_BLUE)
        elif style == "Source Code":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.12)
            set_paragraph_font(paragraph, name="Consolas", size=8.5)
        elif previous_had_graphic and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(6)
            set_paragraph_font(paragraph, size=9.5, italic=True)
        elif has_numbering(paragraph):
            paragraph.style = "List Paragraph"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(3)
            set_paragraph_font(paragraph, size=11)
        elif has_graphic:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(2)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(6)
            set_paragraph_font(paragraph, size=11)

        previous_had_graphic = has_graphic


def flatten_hyperlinks_in_paragraph(paragraph) -> None:
    parent = paragraph._p
    for hyperlink in list(parent.findall(qn("w:hyperlink"))):
        insert_at = parent.index(hyperlink)
        for child in list(hyperlink):
            hyperlink.remove(child)
            parent.insert(insert_at, child)
            insert_at += 1
        parent.remove(hyperlink)


def flatten_hyperlinks(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        flatten_hyperlinks_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    flatten_hyperlinks_in_paragraph(paragraph)


def table_weights(table) -> list[float]:
    cols = len(table.columns)
    header = [cell.text.strip().lower() for cell in table.rows[0].cells] if table.rows else []
    if cols == 2:
        return [2.2, 3.8]
    if cols == 3:
        if header[:3] == ["collection", "index", "purpose"]:
            return [1.35, 1.75, 3.2]
        if header and "helper function" in header[0]:
            return [1.65, 2.75, 2.35]
        return [1.6, 2.0, 2.9]
    if cols == 4:
        return [1.6, 1.45, 2.3, 3.0]
    return [1.0] * cols


def format_tables(doc: Document) -> None:
    content_width = section_content_width_dxa(doc.sections[0])
    for table in doc.tables:
        table.style = "Table Grid"
        widths = column_widths_from_weights(table_weights(table), total_width_dxa=content_width)
        apply_table_geometry(
            table,
            widths,
            table_width_dxa=content_width,
            indent_dxa=0,
            cell_margins_dxa={"top": 95, "bottom": 95, "start": 120, "end": 120},
        )

        if table.rows:
            repeat_header(table.rows[0])
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if r_idx == 0:
                    shade_cell(cell, LIGHT_BLUE)
                for paragraph in cell.paragraphs:
                    paragraph.style = "Normal"
                    paragraph.paragraph_format.space_after = Pt(1)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        set_run_font(run, size=9, bold=(True if r_idx == 0 else None))


def postprocess() -> None:
    doc = Document(PANDOC_OUT)
    configure_document_styles(doc)
    toc_entries = collect_toc_entries(doc)
    flatten_hyperlinks(doc)
    format_paragraphs(doc)
    format_tables(doc)
    insert_front_matter(doc, toc_entries)
    doc.core_properties.title = "OSDP Access Controller Proof of Concept"
    doc.core_properties.subject = "Databases II Course Project"
    doc.core_properties.author = "Martin Velichkovski"
    doc.save(OUTPUT_DOCX)


def main() -> None:
    run_pandoc()
    postprocess()
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
